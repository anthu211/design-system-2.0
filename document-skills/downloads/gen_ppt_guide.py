import shutil, os
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE = "document-skills/document-skill/templates/Prevalent AI - Word Template with Cover Page - 01.08.17.dotx"
OUTPUT = "/Users/ananthusunil/Desktop/Skills for powerpoint/PAI-Claude-PowerPoint-Setup-Guide.docx"


def clear_showcase_section(doc):
    body = doc.element.body
    children = list(body)
    sectPr_para_indices = []
    for i, child in enumerate(children):
        if child.tag == qn('w:p'):
            pPr = child.find(qn('w:pPr'))
            if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                sectPr_para_indices.append(i)
    if len(sectPr_para_indices) >= 2:
        toc_end_idx = sectPr_para_indices[1]
        final_sectPr = children[-1]
        for child in children[toc_end_idx + 1:]:
            if child is not final_sectPr:
                body.remove(child)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = etree.SubElement(tc, qn('w:tcPr'))
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = etree.SubElement(tcPr, qn('w:shd'))
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)


def add_para(doc, text, style):
    doc.add_paragraph(text, style=style)


def add_run_para(doc, text, style):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


def dotx_to_docx(src, dst):
    """Copy a .dotx to .dst and patch the content type so python-docx can open it."""
    import zipfile, re
    with zipfile.ZipFile(src, 'r') as zin:
        with zipfile.ZipFile(dst, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == '[Content_Types].xml':
                    data = data.replace(
                        b'wordprocessingml.template.main+xml',
                        b'wordprocessingml.document.main+xml'
                    )
                zout.writestr(item, data)


def generate():
    tmp = OUTPUT.replace('.docx', '_tmp.docx')
    dotx_to_docx(TEMPLATE, tmp)
    doc = Document(tmp)

    # --- Cover page replacements ---
    replacements = {
        "Document title – 42 pt":   "Claude for PowerPoint — PAI Brand Setup Guide",
        "Document subtitle – 18pt": "How to Create PAI-Branded Presentations Using the Claude Add-In",
        "Month 2017":               "May 2025",
    }
    for para in doc.paragraphs:
        for run in para.runs:
            for old, new in replacements.items():
                if old in run.text:
                    run.text = run.text.replace(old, new)

    # --- Running header ---
    for section in doc.sections:
        for para in section.header.paragraphs:
            for run in para.runs:
                if "Document Title" in run.text:
                    run.text = run.text.replace("Document Title", "Claude for PowerPoint — PAI Brand Setup Guide")

    # --- Clear showcase section BEFORE adding content ---
    clear_showcase_section(doc)

    # =========================================================
    # SECTION 1 — What is this?
    # =========================================================
    add_para(doc, "What is this?", "Heading 1 - 14pt")
    add_para(doc,
        "Prevalent AI has set up a Claude skill inside Microsoft PowerPoint that lets anyone on the team "
        "create a fully branded presentation just by describing what they need. You do not need to be a "
        "designer, you do not need to know the brand colors or fonts, and you do not need to manually "
        "format anything. You type what you want, and Claude builds it — using the correct PAI layouts, "
        "colors, typography, and slide structure every time.",
        "Body Copy - 10pt")

    # =========================================================
    # SECTION 2 — Before You Start
    # =========================================================
    add_para(doc, "Before You Start", "Heading 1 - 14pt")
    add_para(doc,
        "You need three things in place before you can use this for the first time.",
        "Body Copy - 10pt")

    add_para(doc, "1. Microsoft PowerPoint", "Heading 2 - 12pt")
    add_para(doc,
        "This works on PowerPoint for Windows (Microsoft 365), PowerPoint for Mac (version 16.46 or later), "
        "and PowerPoint on the web. If you are unsure which version you have, go to Help → About PowerPoint.",
        "Body Copy - 10pt")

    add_para(doc, "2. The Claude Add-In Installed in PowerPoint", "Heading 2 - 12pt")
    add_para(doc,
        "The Claude add-in adds a Claude panel to your PowerPoint toolbar. If you do not see a Claude button "
        "in your Home tab, the add-in has not been installed on your machine yet. Contact your IT admin and "
        "ask them to install it, or follow the installation steps in the Admin Setup section at the end of "
        "this guide.",
        "Body Copy - 10pt")

    add_para(doc, "3. The PAI PowerPoint Template File", "Heading 2 - 12pt")
    add_para(doc,
        "This is the official Prevalent AI template that contains the correct brand colors, fonts, and slide "
        "layouts. It is stored in the shared templates folder. Download it and keep it somewhere easy to find "
        "— you will need to open it every time you start a new deck.",
        "Body Copy - 10pt")

    # =========================================================
    # SECTION 3 — How It Works
    # =========================================================
    add_para(doc, "How It Works", "Heading 1 - 14pt")
    add_para(doc,
        "The Claude add-in reads the design settings of whichever PowerPoint file you currently have open. "
        "This means that when you open the PAI template and then use Claude, it automatically knows the "
        "Prevalent AI colors, fonts, and layouts — and uses them when it builds your slides.",
        "Body Copy - 10pt")
    add_para(doc,
        "On top of that, Prevalent AI has set up a custom skill called PAI Pitch Deck Generator that is "
        "available inside the Claude panel. When you select this skill, Claude also loads a set of brand "
        "rules specific to Prevalent AI — things like how to structure a deck, what each slide type should "
        "look like, what never to put on a cover slide, and how many bullet points are allowed per slide. "
        "You never need to explain any of this to Claude yourself. You simply select the skill, describe the "
        "presentation you need, and Claude handles everything else.",
        "Body Copy - 10pt")

    # =========================================================
    # SECTION 4 — Creating a Presentation: Step by Step
    # =========================================================
    add_para(doc, "Creating a Presentation — Step by Step", "Heading 1 - 14pt")
    add_para(doc,
        "Follow these steps every time you need to build a new deck. No design experience is required — "
        "Claude handles all formatting and branding automatically.",
        "Body Copy - 10pt")

    # Step 1
    add_para(doc, "Step 1 — Open the PAI Template", "Heading 2 - 12pt")
    add_para(doc,
        "The most important thing to remember is that you must always start from the PAI template file, not "
        "from a blank PowerPoint. The template is what gives Claude the correct brand colors and layouts to "
        "work with.",
        "Body Copy - 10pt")
    add_para(doc,
        "Go to the shared templates folder, find the file called Template_PAI_Presentation.pptx, and open "
        "it in PowerPoint. The file will open with the PAI branding already applied. You do not need to "
        "change anything — just leave it open and move to the next step.",
        "Body Copy - 10pt")
    add_para(doc,
        "If you try to use Claude in a blank or non-PAI file, it will not know to use the Prevalent AI "
        "brand and the output will not look correct.",
        "Body Copy - 10pt")

    # Step 2
    add_para(doc, "Step 2 — Open the Claude Panel", "Heading 2 - 12pt")
    add_para(doc,
        "Once your PAI template is open in PowerPoint, look at the ribbon across the top of the screen and "
        "click the Home tab. You should see a Claude button in the toolbar. Click it. A panel will slide "
        "open on the right side of your screen — this is where you will interact with Claude.",
        "Body Copy - 10pt")
    add_para(doc,
        "If you do not see the Claude button in the toolbar, the add-in is not installed on your machine. "
        "Contact your IT admin and ask them to install the Claude for PowerPoint add-in.",
        "Body Copy - 10pt")

    # Step 3
    add_para(doc, "Step 3 — Select the PAI Pitch Deck Generator Skill", "Heading 2 - 12pt")
    add_para(doc,
        "At the top of the Claude panel, you will see a Skills dropdown or selector. Click on it. A list of "
        "available skills will appear. Find and select PAI Pitch Deck Generator.",
        "Body Copy - 10pt")
    add_para(doc,
        "Once selected, this skill is active for your session. It loads all of the Prevalent AI brand rules "
        "into Claude automatically — the correct slide order, content limits, what each slide type should "
        "contain, and what to avoid. You do not need to explain any of this yourself.",
        "Body Copy - 10pt")

    # Step 4
    add_para(doc, "Step 4 — Describe What You Need", "Heading 2 - 12pt")
    add_para(doc,
        "In the text box at the bottom of the Claude panel, type a description of the presentation you want "
        "in plain English. You do not need to use any special format or technical language — just write "
        "naturally, as you would describe it to a colleague.",
        "Body Copy - 10pt")
    add_para(doc, "You can be as brief or as detailed as you like. Here are some examples:", "Body Copy - 10pt")

    add_para(doc,
        "\"Create a presentation on our Q3 vendor risk assessment results. The audience is executive "
        "leadership. Cover the key findings, top risks identified, current remediation status, and "
        "recommended next steps. About 10 to 12 slides.\"",
        "Bulleted List - 10pt")
    add_para(doc,
        "\"Build a board presentation titled 2025 Vendor Risk Program Update. Include four sections: "
        "Program Overview, Key Findings, Remediation Progress, and Recommendations. Keep it concise — "
        "around 12 slides total.\"",
        "Bulleted List - 10pt")
    add_para(doc,
        "\"Add an agenda slide listing these five topics: Introduction, Threat Landscape, Program Status, "
        "Recommendations, and Q&A.\"",
        "Bulleted List - 10pt")
    add_para(doc,
        "\"Add a new section on supply chain breach statistics. Include three content slides covering recent "
        "industry incidents, common entry points, and what we are doing to address this.\"",
        "Bulleted List - 10pt")

    add_para(doc,
        "Once you have typed your request, press Enter or click the Send button. Claude will begin building "
        "the slides directly in your open PowerPoint file. Depending on the number of slides requested, this "
        "usually takes between 15 and 60 seconds.",
        "Body Copy - 10pt")

    # Step 5
    add_para(doc, "Step 5 — Review Your Slides", "Heading 2 - 12pt")
    add_para(doc,
        "Once Claude has finished generating the slides, scroll through the deck in PowerPoint and review "
        "the content. Check that the information is accurate, the key points are clear, and the structure "
        "makes sense for your audience.",
        "Body Copy - 10pt")
    add_para(doc,
        "If you want to change the wording on a slide, simply click on the text in PowerPoint and edit it "
        "as you normally would. If you want Claude to make a larger change — such as adding a slide, "
        "removing a section, or rewriting a slide — you can ask it directly in the chat panel. For example:",
        "Body Copy - 10pt")
    add_para(doc, "\"Make slide 4 shorter — reduce it to three bullet points.\"", "Bulleted List - 10pt")
    add_para(doc, "\"Add one more slide after slide 6 covering remediation timelines.\"", "Bulleted List - 10pt")
    add_para(doc, "\"Rewrite the recommendations slide to be more direct and action-oriented.\"", "Bulleted List - 10pt")

    # Step 6
    add_para(doc, "Step 6 — Save Your Presentation", "Heading 2 - 12pt")
    add_para(doc,
        "When you are happy with the deck, go to File → Save As and save it with a clear, descriptive "
        "filename — for example: Q3 Vendor Risk Assessment — Board Update — May 2025.pptx. Save it to your "
        "preferred location: your local drive, SharePoint, or OneDrive.",
        "Body Copy - 10pt")
    add_para(doc,
        "Do not save over the original Template_PAI_Presentation.pptx file. Always use Save As and give "
        "your new deck a different name.",
        "Body Copy - 10pt")

    # =========================================================
    # SECTION 5 — What Claude Always Produces
    # =========================================================
    add_para(doc, "What Claude Always Produces Automatically", "Heading 1 - 14pt")
    add_para(doc,
        "When you use the PAI Pitch Deck Generator skill, the following elements are handled for you every "
        "time — you never need to ask for them:",
        "Body Copy - 10pt")

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Prevalent AI Table Style'
    hdr = table.rows[0].cells
    hdr[0].paragraphs[0].text = ""
    hdr[1].paragraphs[0].text = ""
    for cell in hdr:
        set_cell_bg(cell, '372355')
        cell.paragraphs[0].style = doc.styles['Table Heading - 10pt']
    hdr[0].paragraphs[0].runs[0].text = "Element"
    hdr[1].paragraphs[0].runs[0].text = "What you get"

    rows_data = [
        ("Cover slide",
         "Slide 1 — PAI logo only, dark purple gradient background. No text is placed on it. This is intentional and correct."),
        ("Title slide",
         "Slide 2 — Dark purple background with your deck title in large white text."),
        ("Section dividers",
         "Branded divider slides in dark purple or orange, automatically inserted between major sections."),
        ("Content slides",
         "Clean white background, PAI purple title at the top, and bullet points in the correct brand style."),
        ("Back cover",
         "Last slide — PAI logo only, matching the cover slide. No text."),
        ("Slide numbers",
         "Every interior slide includes a slide number in the bottom-right corner in the correct PAI format."),
        ("Colors",
         "PAI brand palette throughout. Office default colors are never used."),
        ("Fonts",
         "Calibri Light for headings and Calibri for body text throughout the entire deck."),
    ]
    for label, desc in rows_data:
        row = table.add_row().cells
        row[0].paragraphs[0].text = ""
        row[1].paragraphs[0].text = ""
        row[0].paragraphs[0].style = doc.styles['Body Copy - 10pt']
        row[1].paragraphs[0].style = doc.styles['Body Copy - 10pt']
        p0 = row[0].paragraphs[0]
        r0 = p0.add_run(label)
        r0.bold = True
        row[1].paragraphs[0].add_run(desc)

    doc.add_paragraph("", "Body Copy - 10pt")

    # =========================================================
    # SECTION 6 — Troubleshooting
    # =========================================================
    add_para(doc, "Troubleshooting", "Heading 1 - 14pt")

    issues = [
        (
            "The slides look plain or use the wrong colors.",
            "You opened Claude in a blank or non-PAI file. Close the presentation without saving, go back to "
            "the shared templates folder, open Template_PAI_Presentation.pptx, and try again."
        ),
        (
            "There is text on the cover slide.",
            "The PAI Pitch Deck Generator skill was not selected before generating. Go to the Skills selector "
            "in the Claude panel, select PAI Pitch Deck Generator, and ask Claude to fix the cover slide."
        ),
        (
            "Slide 2 has a white background instead of dark purple.",
            "Ask Claude directly in the chat: \"Please change slide 2 to a dark purple title slide.\" "
            "Claude will correct it immediately."
        ),
        (
            "There are no slide numbers on the slides.",
            "Ask Claude in the chat: \"Please add slide numbers to all slides except the cover and back "
            "cover.\" Claude will add them in the correct format."
        ),
        (
            "I cannot see the Claude button in PowerPoint.",
            "The add-in is not installed on your machine. Contact your IT admin and ask them to install the "
            "Claude for PowerPoint add-in from the Microsoft 365 admin centre."
        ),
    ]

    for problem, fix in issues:
        add_para(doc, problem, "Heading 3 - 10pt")
        add_para(doc, fix, "Body Copy - 10pt")

    # =========================================================
    # SECTION 7 — Admin Setup
    # =========================================================
    add_para(doc, "Admin Setup — IT Only", "Heading 1 - 14pt")
    add_para(doc,
        "This section is for IT admins setting up the Claude add-in and skill for the first time. "
        "End users do not need to read this section.",
        "Body Copy - 10pt")

    add_para(doc, "Installing the Add-In", "Heading 2 - 12pt")
    add_para(doc,
        "The Claude for PowerPoint add-in can be deployed org-wide from the Microsoft 365 admin centre, "
        "so that all users receive it automatically without needing to install it themselves. Alternatively, "
        "individual users can install it by opening PowerPoint, going to Insert → Add-ins → Get Add-ins, "
        "searching for Claude, and clicking Add.",
        "Body Copy - 10pt")

    add_para(doc, "Adding the PAI Pitch Deck Generator Skill", "Heading 2 - 12pt")
    add_para(doc,
        "The PAI Pitch Deck Generator skill must be added to your Claude organisation so that it appears in "
        "the Skills list for all users. Log in to claude.ai using an org admin account, go to Settings → "
        "Organisation → Skills, upload the file pai-ppt-org-instructions.md, and save. Once added, the skill "
        "will appear in the Skills dropdown inside the Claude panel in PowerPoint for every user in your "
        "organisation.",
        "Body Copy - 10pt")

    add_para(doc, "Sharing the Template File", "Heading 2 - 12pt")
    add_para(doc,
        "Upload Template_PAI_Presentation.pptx to your organisation's shared drive — SharePoint, OneDrive, "
        "or Google Drive — and share the link with all staff. Include the link on your internal wiki or "
        "intranet so it is easy for anyone to find.",
        "Body Copy - 10pt")

    doc.save(OUTPUT)
    os.remove(tmp)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    generate()
